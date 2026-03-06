from docx import Document
from docx.shared import Inches, Pt

def generate_proposal_docx(content: str, output_path: str):
    """Generates a Word document from proposal content."""
    doc = Document()
    
    # Simplified markdown parsing
    sections = content.split('\n')
    for section in sections:
        if section.startswith('# '):
            doc.add_heading(section[2:], 0)
        elif section.startswith('## '):
            doc.add_heading(section[3:], 1)
        elif section.strip():
            doc.add_paragraph(section)
            
    doc.save(output_path)
